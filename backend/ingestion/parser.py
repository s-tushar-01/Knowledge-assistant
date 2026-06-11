import logging
from pathlib import Path
from typing import List

from llama_index.core import Document

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {
    ".pdf", ".docx", ".doc", ".txt", ".md",
}


def parse_document(file_path: str | Path, document_id: str) -> List[Document]:
    """
    Parse a file with Unstructured and return LlamaIndex Document objects.
    Each logical section becomes one Document so the chunker can split them further.
    """
    file_path = Path(file_path)
    ext = file_path.suffix.lower()

    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file extension: {ext}")

    logger.info(f"Parsing {file_path.name} ({ext})")

    if ext == ".docx":
        return _parse_docx(file_path, document_id)

    if ext in {".txt", ".md"}:
        text = file_path.read_text(encoding="utf-8", errors="ignore").strip()
        return [_make_doc(text, file_path, document_id, 1, "")] if text else []

    try:
        from unstructured.partition.auto import partition
        elements = partition(filename=str(file_path))
    except Exception as exc:
        logger.error(f"Unstructured failed for {file_path}: {exc}")
        raise

    documents: List[Document] = []
    current_section = ""
    current_page = 1
    buffer: List[str] = []

    def flush():
        if buffer:
            documents.append(_make_doc(
                "\n".join(buffer), file_path, document_id,
                current_page, current_section,
            ))
            buffer.clear()

    for el in elements:
        el_type = type(el).__name__
        text = str(el).strip()
        if not text:
            continue

        # Track page numbers from Unstructured metadata
        if hasattr(el, "metadata") and hasattr(el.metadata, "page_number"):
            pg = el.metadata.page_number
            if pg:
                current_page = pg

        # Section headings → flush current buffer, start new section
        if el_type in ("Title", "Header"):
            flush()
            current_section = text

        buffer.append(text)

    flush()

    # Fallback if Unstructured produced no structure
    if not documents:
        raw = "\n".join(str(e) for e in elements if str(e).strip())
        if raw:
            documents.append(_make_doc(raw, file_path, document_id, 1, ""))

    logger.info(
        f"Parsed {file_path.name}: {len(documents)} sections, "
        f"{sum(len(d.text) for d in documents)} chars total"
    )
    return documents


def _parse_docx(file_path: Path, document_id: str) -> List[Document]:
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:
        raise RuntimeError("DOCX support requires python-docx. Add python-docx to requirements.txt.") from exc

    docx = DocxDocument(str(file_path))
    documents: List[Document] = []
    current_section = ""
    buffer: List[str] = []

    def flush():
        if buffer:
            documents.append(_make_doc(
                "\n".join(buffer).strip(), file_path, document_id, 1, current_section,
            ))
            buffer.clear()

    for paragraph in docx.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        style_name = paragraph.style.name.lower() if paragraph.style and paragraph.style.name else ""
        if style_name.startswith("heading"):
            flush()
            current_section = text

        buffer.append(text)

    flush()

    table_text: List[str] = []
    for table in docx.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                table_text.append(" | ".join(cells))

    if table_text:
        documents.append(_make_doc("\n".join(table_text), file_path, document_id, 1, "Tables"))

    logger.info(
        "Parsed %s with python-docx: %d sections, %d chars total",
        file_path.name, len(documents), sum(len(d.text) for d in documents),
    )
    return documents


def _make_doc(text: str, file_path: Path, document_id: str, page: int, section: str) -> Document:
    return Document(
        text=text,
        metadata={
            "document_id": document_id,
            "source_file": file_path.name,
            "file_path": str(file_path),
            "source_type": file_path.suffix.lstrip(".").lower(),
            "page_number": page,
            "section_heading": section,
        },
        excluded_llm_metadata_keys=["file_path", "document_id"],
        excluded_embed_metadata_keys=["file_path", "document_id"],
    )
